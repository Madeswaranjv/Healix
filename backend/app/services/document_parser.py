"""Document parsing and text chunking service for PDF, DOCX, TXT, CSV, JSON, and MD files."""
import io
import logging
from pathlib import Path
from typing import List, Tuple
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes using a multi-engine fallback strategy:
    1. pdfplumber (best for layout and tables)
    2. pypdf (fast, robust for text streams)
    3. pypdfium2 (high-fidelity Chromium PDF engine)
    4. pdfminer.six (low-level stream extractor)
    """
    extracted_text = ""

    # Strategy 1: pdfplumber
    try:
        import pdfplumber
        text_pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_pages.append(page_text.strip())
        if text_pages:
            extracted_text = "\n\n".join(text_pages).strip()
            if len(extracted_text) > 20:
                return extracted_text
    except Exception as e:
        logger.debug(f"pdfplumber extraction failed or returned empty: {e}")

    # Strategy 2: pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_pages.append(page_text.strip())
        if text_pages:
            extracted_text = "\n\n".join(text_pages).strip()
            if len(extracted_text) > 20:
                return extracted_text
    except Exception as e:
        logger.debug(f"pypdf extraction failed or returned empty: {e}")

    # Strategy 3: pypdfium2
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(io.BytesIO(file_bytes))
        text_pages = []
        for page in pdf:
            textpage = page.get_textpage()
            page_text = textpage.get_text_range()
            if page_text and page_text.strip():
                text_pages.append(page_text.strip())
        if text_pages:
            extracted_text = "\n\n".join(text_pages).strip()
            if len(extracted_text) > 20:
                return extracted_text
    except Exception as e:
        logger.debug(f"pypdfium2 extraction failed or returned empty: {e}")

    # Strategy 4: pdfminer.six
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        mined_text = pdfminer_extract(io.BytesIO(file_bytes))
        if mined_text and mined_text.strip():
            return mined_text.strip()
    except Exception as e:
        logger.debug(f"pdfminer extraction failed: {e}")

    return extracted_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes including paragraphs, headings, and tables."""
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    return "\n".join(full_text).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extracts text from plain text, markdown, CSV, or JSON bytes with auto-encoding detection."""
    for encoding in ["utf-8", "utf-8-sig", "latin-1", "utf-16", "cp1252"]:
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def parse_document(filename: str, file_bytes: bytes) -> str:
    """Parses a document into plain text according to its file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext in [".txt", ".md", ".csv", ".json", ".tsv", ".rtf"]:
        return extract_text_from_txt(file_bytes)
    else:
        # Default fallback
        return extract_text_from_txt(file_bytes)


def chunk_text(
    text: str,
    chunk_size: int = 800,
    chunk_overlap: int = 150
) -> List[str]:
    """Splits raw text into manageable chunks using RecursiveCharacterTextSplitter."""
    if not text or not text.strip():
        return []
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", "; ", ", ", " ", ""]
    )
    return splitter.split_text(text.strip())


def process_uploaded_document(
    filename: str,
    file_bytes: bytes
) -> Tuple[str, List[str]]:
    """Extracts text and chunks an uploaded document. Returns (full_text, chunks)."""
    full_text = parse_document(filename, file_bytes)
    if not full_text:
        return "", []
    chunks = chunk_text(full_text)
    return full_text, chunks
